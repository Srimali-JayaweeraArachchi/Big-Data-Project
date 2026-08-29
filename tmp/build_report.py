from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\srimali\Desktop\Semester 8\EC8202-Big Data Analytics\Big Data Project")
SHOTS = ROOT / "docs" / "screenshots"
OUT = ROOT / "output" / "reports"
TMP = ROOT / "tmp" / "report_assets"
OUT.mkdir(parents=True, exist_ok=True)
TMP.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "EC8203_Smart_Grid_Technical_Report.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT = "EAF2F8"
GRAY = "5B6573"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_picture(doc, filename, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(SHOTS / filename), width=Inches(width))
    add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8.5)
    return table


def add_section_title(doc, number, title, lead=None):
    doc.add_heading(f"{number}. {title}", level=1)
    if lead:
        p = doc.add_paragraph(lead)
        p.style = doc.styles["Quote"]


def page_break(doc):
    doc.add_page_break()


def make_architecture():
    path = TMP / "architecture.png"
    img = Image.new("RGB", (1800, 950), "white")
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 35)
        small = ImageFont.truetype("arial.ttf", 27)
        title = ImageFont.truetype("arialbd.ttf", 43)
    except OSError:
        font = small = title = ImageFont.load_default()
    d.text((55, 35), "Smart Grid Kappa-Oriented Architecture", font=title, fill="#17365D")
    boxes = {
        "Meter Simulator\nPython / JSON": (70, 170, 370, 310),
        "Apache Kafka\n3 partitions": (470, 170, 770, 310),
        "Spark Structured\nStreaming": (870, 170, 1210, 310),
        "PostgreSQL\nQueryable store": (1370, 320, 1700, 470),
        "Tariff Generator\nCSV every 5 min": (70, 570, 370, 710),
        "Shared Batch\nVolume": (470, 570, 770, 710),
        "Apache Airflow\nValidate + bill": (870, 570, 1210, 710),
        "FastAPI\nServing + health": (1090, 780, 1390, 900),
        "Streamlit\nDashboard": (1450, 780, 1730, 900),
    }
    for label, box in boxes.items():
        d.rounded_rectangle(box, radius=20, fill="#EAF2F8", outline="#2E74B5", width=4)
        lines = label.split("\n")
        for j, line in enumerate(lines):
            bb = d.textbbox((0, 0), line, font=font if j == 0 else small)
            x = (box[0] + box[2] - (bb[2]-bb[0])) / 2
            y = box[1] + 28 + j * 48
            d.text((x, y), line, font=font if j == 0 else small, fill="#17365D")
    def arrow(a, b):
        d.line([a, b], fill="#F28C28", width=8)
        x, y = b
        d.polygon([(x, y), (x-22, y-14), (x-22, y+14)], fill="#F28C28")
    arrow((370,240),(470,240)); arrow((770,240),(870,240)); arrow((1210,240),(1370,390))
    arrow((370,640),(470,640)); arrow((770,640),(870,640)); arrow((1210,640),(1370,450))
    arrow((1535,470),(1240,780)); arrow((1390,840),(1450,840))
    d.text((75, 330), "Continuous streaming path", font=small, fill="#5B6573")
    d.text((75, 730), "Simulated daily batch path", font=small, fill="#5B6573")
    img.save(path)
    return path


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(0.75)
section.left_margin = section.right_margin = Inches(0.82)
section.header_distance = section.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(35, 43, 52)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08
for name, size, before, after, color in [
    ("Heading 1", 16, 12, 6, BLUE), ("Heading 2", 13, 9, 4, BLUE),
    ("Heading 3", 11.5, 7, 3, NAVY)]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True
for list_name in ("List Bullet", "List Number"):
    s = styles[list_name]
    s.font.name = "Calibri"; s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.5)
    s.paragraph_format.first_line_indent = Inches(-0.25)
    s.paragraph_format.space_after = Pt(3)
styles["Quote"].font.name = "Calibri"
styles["Quote"].font.size = Pt(10)
styles["Quote"].font.italic = True
styles["Quote"].font.color.rgb = RGBColor.from_string(GRAY)

header = section.header.paragraphs[0]
header.text = "EC8203 Applied Big Data Engineering | Technical Report"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
add_page_number(section.footer.paragraphs[0])
for r in section.footer.paragraphs[0].runs:
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(105); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TECHNICAL REPORT"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string("F28C28")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(10)
r = p.add_run("Smart Grid Energy\nMonitoring & Billing"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(38)
r = p.add_run("A Kappa-Oriented End-to-End Big Data Pipeline"); r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(GRAY)
add_table(doc, ["Module", "Assessment"], [["EC8203 Applied Big Data Engineering", "Mini Project - 25%"]], [4680,4680])
doc.add_paragraph()
for label, value in [("Student Name", "[Enter student name]"), ("Registration Number", "[Enter registration number]"), ("Submission Type", "Individual / Group [edit as applicable]"), ("Date", "29 August 2026")]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a=p.add_run(label+": "); a.bold=True; b=p.add_run(value)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(38)
r=p.add_run("Implemented with Apache Kafka, Spark Structured Streaming, Apache Airflow, PostgreSQL, FastAPI, Streamlit and Docker Compose"); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(GRAY)

page_break(doc)
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph("This project implements a small but complete data platform for a utility company that needs live visibility into electricity demand and solar contribution while producing household bills from a periodic tariff feed. The solution uses a Kappa-oriented architecture: smart-meter events enter a single Kafka stream and are processed continuously by Spark Structured Streaming. A deliberately small Airflow-managed batch path handles external tariff reference files and materialises a consolidated billing report in PostgreSQL.")
doc.add_paragraph("The demonstrator simulates 30 households in four grid zones. Meter events are emitted every two seconds and a simulated day is compressed to five minutes. Processing includes schema validation, quarantine of invalid events, event-time handling, net-grid and renewable-ratio enrichment, one-minute zone aggregation, daily household aggregation, tariff reconciliation, subsidy-aware billing, and reproducible threshold alerts. FastAPI exposes query and health endpoints, while Streamlit presents a live dashboard.")
doc.add_paragraph("End-to-end verification showed all eight Docker services running, successful Spark micro-batches, successful Airflow DAG executions, 30 tariff records, 30 household billing records, and a healthy serving API. The implementation satisfies the assignment requirements while remaining reproducible on a single Windows laptop.")
doc.add_heading("Report Map", level=2)
add_numbered(doc, ["Business problem and requirements", "Architecture decision: Lambda versus Kappa", "Architecture and technology stack", "Ingestion, processing, storage and serving", "Observability, results and testing", "Limitations, production evolution and conclusion"])
doc.add_heading("Key Assumptions", level=2)
add_bullets(doc, ["One simulated day equals five minutes.", "Meter values represent interval kWh rather than cumulative counters.", "Thirty households are distributed across EAST, NORTH, SOUTH and WEST zones.", "A 15% demonstration discount is applied where subsidy_flag is true.", "The deployment is an educational single-node environment, not a production high-availability cluster."])

page_break(doc)
add_section_title(doc, "1", "Use Case and Business Requirements", "Chosen use case: Smart Grid Energy Monitoring & Billing (Use Case 3).")
doc.add_paragraph("The utility needs two views of the same operating reality. Operations staff need near-real-time information about load and renewable contribution by zone. Billing staff need a periodic reconciliation that combines household usage with the latest valid tariff and subsidy attributes. The consolidated system therefore has to combine a continuous event source with an independently arriving reference-data file.")
doc.add_heading("1.1 Business Question", level=2)
doc.add_paragraph("What is the current grid load and renewable contribution by zone, and what will each household's bill look like once daily tariff data is applied to consumption?")
doc.add_heading("1.2 Functional Requirements", level=2)
add_bullets(doc, ["Generate continuous smart-meter events containing meter, household, consumption, solar, zone and event timestamp fields.", "Generate one tariff CSV per simulated day with household, rate, tier, subsidy and effective-date fields.", "Clean, enrich, aggregate and join the two sources rather than merely forwarding records.", "Persist results in a queryable store and expose a consolidated dashboard and API.", "Provide structured logs and at least one health or alert rule."])
doc.add_heading("1.3 Non-Functional Requirements", level=2)
add_table(doc, ["Requirement", "Interpretation"], [
    ("Latency", "Seconds-to-minutes for operational metrics; five-minute simulated daily billing."),
    ("Replay", "Kafka retention and Spark checkpoints allow stream recovery without a second code path."),
    ("Reliability", "Producer acknowledgements, retries, idempotent keys/upserts and Airflow retries."),
    ("Observability", "JSON logs, database heartbeats, API health, metrics and stored alerts."),
    ("Reproducibility", "Docker Compose starts the complete platform with documented commands."),
], [2200,7160])
doc.add_heading("1.4 Data Contracts", level=2)
doc.add_paragraph("Streaming contract: meter_id, household_id, power_consumption_kwh, solar_generation_kwh, grid_zone and timestamp. Batch contract: household_id, tariff_rate, billing_tier, subsidy_flag and effective_date. Invalid stream events are quarantined; invalid tariff files fail visibly and are retryable.")

page_break(doc)
add_section_title(doc, "2", "Architecture Decision: Lambda vs Kappa")
doc.add_paragraph("A Lambda architecture would maintain a low-latency speed layer and an independent batch layer over historical meter data. A Kappa architecture uses the event log as the main source of truth and applies one streaming transformation path; historical correction is achieved by replaying retained events. The chosen design is Kappa-oriented because meter events dominate the workload and the two-week implementation benefits from avoiding duplicated business logic.")
add_table(doc, ["Criterion", "Lambda", "Kappa-oriented choice"], [
    ("Latency", "Strong via speed layer", "Strong via continuous Spark processing"),
    ("Replay", "Batch recomputation is explicit", "Kafka replay plus checkpoints"),
    ("Consistency", "Two implementations can diverge", "One meter-event transformation path"),
    ("Cost", "More compute and storage layers", "Lower single-node footprint"),
    ("Complexity", "Batch and speed code must be maintained", "Simpler operational model"),
    ("Reference feed", "Fits batch layer", "Retained as a small Airflow-managed side path"),
], [1650,3650,4060])
doc.add_heading("2.1 Explicit Justification", level=2)
doc.add_paragraph("The live business question requires low latency, and Spark Structured Streaming already provides windowed aggregation and fault-tolerant checkpoints. Maintaining a second implementation for historical meter processing would increase code volume and create a consistency risk without adding proportional value at this scale. Kafka provides a durable boundary between production and processing and supports controlled replay. PostgreSQL primary keys and upserts make sink effects idempotent during repeated micro-batches.")
doc.add_heading("2.2 Rejected Alternative and Honest Trade-off", level=2)
doc.add_paragraph("Lambda was rejected because separate speed and batch calculations for consumption, solar ratio and net grid load would duplicate logic. However, Kappa is not free of trade-offs: a long replay can compete with the live Spark job, retention must be sized correctly, and correcting historical reference data may require deliberate reprocessing. The Airflow tariff path is not a second meter-data computation layer; it is a scheduled reference-data workflow required by the source's delivery pattern.")

page_break(doc)
add_section_title(doc, "3", "System Architecture and Data Flow")
arch = make_architecture()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(arch), width=Inches(6.5))
add_caption(doc, "Figure 1. Kappa-oriented architecture with continuous meter and simulated-daily tariff paths.")
doc.add_heading("3.1 Streaming Path", level=2)
add_numbered(doc, ["The Python producer publishes keyed JSON meter events to a three-partition Kafka topic.", "Spark parses each event, validates required fields and timestamp, enriches the record and writes micro-batch outputs.", "PostgreSQL stores clean readings, rejected events, zone windows, household usage, heartbeats and alerts.", "FastAPI serves database results and Streamlit refreshes the operational view every ten seconds."])
doc.add_heading("3.2 Batch Path", level=2)
add_numbered(doc, ["The tariff generator writes a temporary file and atomically renames it to a CSV every five minutes.", "Airflow discovers unprocessed files, validates their schema and values, and performs idempotent tariff upserts.", "The DAG joins daily household usage to the latest applicable tariff and refreshes the billing report.", "Processed files are archived and audit/heartbeat records make each run diagnosable."])
doc.add_heading("3.3 Deployment", level=2)
doc.add_paragraph("Docker Compose defines eight services and named volumes for Kafka, PostgreSQL, Spark checkpoints, Airflow metadata and shared batch files. Host ports are 5432 (PostgreSQL), 8000 (API), 8081 (Airflow), 8501 (dashboard) and 9092 (Kafka). Service health and dependency conditions reduce startup races.")

add_section_title(doc, "4", "Technology Stack Selection")
add_table(doc, ["Layer", "Technology", "Use-case justification"], [
    ("Simulation", "Python 3.12", "Simple, controllable event/file generation with reproducible anomalies."),
    ("Ingestion", "Apache Kafka 4.1.1", "Durable partitioned buffer, producer/consumer decoupling and replay."),
    ("Streaming", "Spark 3.5.9", "Structured Streaming supports event-time transformations, windows and checkpoints."),
    ("Orchestration", "Airflow 3.1", "Visible, scheduled, retryable tariff ingestion and billing workflow."),
    ("Storage", "PostgreSQL 16", "Queryable relational joins, indexes and upserts match billing and serving needs."),
    ("Serving", "FastAPI", "Lightweight typed API plus health and Prometheus-format metrics."),
    ("Dashboard", "Streamlit", "Rapid operational visualisation directly from serving endpoints."),
    ("Deployment", "Docker Compose", "Repeatable multi-service setup suitable for assessment and live demo."),
], [1350,2150,5860])
doc.add_heading("4.1 Why PostgreSQL Instead of Cassandra or HDFS", level=2)
doc.add_paragraph("The serving workload is dominated by indexed household lookups, date-based reports and joins between usage and tariff data. PostgreSQL fits those relational operations and enables a compact demonstrator. Cassandra would be valuable for very high write throughput but would require query-first denormalisation; HDFS/Parquet would support historical analytics but would add a serving engine for interactive API queries.")
doc.add_heading("4.2 Configuration and Security Scope", level=2)
doc.add_paragraph("Runtime values are supplied through .env variables with documented development defaults. Containers run as non-root where practical, and persistent data is held in named volumes. The report does not claim production security: development credentials, plaintext internal traffic and unauthenticated serving endpoints must be replaced by secrets management, TLS and role-based access control in production.")
add_picture(doc, "01-docker-services-running.png", "Figure 2. Docker Desktop showing the complete multi-service platform running.", 6.15)

add_section_title(doc, "5", "Data Ingestion and Processing Implementation")
doc.add_heading("5.1 Continuous Source", level=2)
doc.add_paragraph("The meter producer cycles through 30 household/meter identities and four zones, emits one event every two seconds, keys Kafka messages by meter_id and requests acks=all with retries. Reproducible anomalies are injected: every 18th event has zero solar and every 25th event has high consumption. This makes the alert path demonstrable without waiting for random rare events.")
doc.add_heading("5.2 Simulated-Daily Source", level=2)
doc.add_paragraph("The tariff producer generates 30 rows every five minutes. It writes to a .tmp file first and renames only after completion, preventing Airflow from reading a partial file. The Airflow DAG validates required columns, positive rates and household identifiers, then upserts on household_id and effective_date. Audit status and archived files make retries idempotent and observable.")
doc.add_heading("5.3 Stream Transformations", level=2)
add_bullets(doc, ["Parse JSON into an explicit schema and convert event timestamps to UTC.", "Reject missing identifiers, invalid timestamps and negative consumption/solar values.", "Calculate net_grid_usage_kwh = consumption - solar.", "Calculate renewable_ratio = solar / consumption, with zero-consumption protection.", "Aggregate one-minute event-time windows by grid zone.", "Aggregate daily usage per household for later tariff reconciliation.", "Create LOW_RENEWABLE warnings below 0.10 and HIGH_GRID_LOAD critical alerts above 8.0 kWh."])
doc.add_heading("5.4 Billing Transformation", level=2)
doc.add_paragraph("For each household-day, Airflow selects the latest tariff whose effective date is not later than the usage date. Billable grid usage is max(net grid usage, 0), so exported solar does not create a negative bill. The demonstration equation is:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("estimated_bill = billable_grid_kwh x tariff_rate x (0.85 if subsidised, otherwise 1.00)"); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
add_picture(doc, "10-structured-producer-logs.png", "Figure 3. Structured JSON logs confirm continuous events entering the ingestion layer.", 6.1)

add_section_title(doc, "6", "Storage, Serving and Dashboard")
doc.add_heading("6.1 PostgreSQL Data Model", level=2)
add_table(doc, ["Table", "Purpose / key"], [
    ("meter_readings_clean", "Validated event-level facts; event identifier supports idempotency."),
    ("rejected_events", "Quarantined payloads and validation reasons."),
    ("zone_metrics", "One-minute zone aggregations for operational charts."),
    ("household_usage_daily", "Daily consumption, solar and net-grid usage per household."),
    ("tariffs", "Effective-dated household tariff and subsidy reference."),
    ("daily_billing_report", "Consolidated household-day billing result."),
    ("pipeline_alerts", "Threshold breaches with severity, metric and threshold."),
    ("pipeline_heartbeats", "Latest service status and structured details."),
], [2650,6710])
doc.add_heading("6.2 Serving API", level=2)
doc.add_paragraph("FastAPI separates machine-readable serving from presentation. Endpoints expose health, overview, zones, time series, alerts, billing, pipeline status and Prometheus-format metrics. SQL queries are deliberately read-only and indexed around common time and household access paths.")
doc.add_heading("6.3 Dashboard", level=2)
doc.add_paragraph("Streamlit polls the API every ten seconds and presents key indicators, zone comparisons, recent consumption versus solar, alert details, the daily billing table and pipeline health. This directly answers both parts of the business question without requiring database access.")
add_picture(doc, "02-dashboard-overview.png", "Figure 4. Live dashboard overview with grid load, renewable contribution and zone comparison.", 6.2)

add_section_title(doc, "7", "Observability Design")
doc.add_paragraph("Observability was designed as part of the pipeline rather than added only to the dashboard. Each stage produces evidence that helps distinguish an ingestion failure, processing backlog, storage error or serving problem.")
add_table(doc, ["Signal", "Implementation", "Diagnostic value"], [
    ("Structured logs", "JSON service, timestamp, level, event and context fields", "Searchable evidence of publish and micro-batch activity"),
    ("Heartbeats", "Spark and Airflow rows in pipeline_heartbeats", "Shows last successful work and component status"),
    ("Health rule", "API unhealthy if latest event age exceeds 120 seconds", "Detects stopped producer or broken stream"),
    ("Metrics", "/metrics exposes event age and unresolved alerts", "Supports external scraping and thresholding"),
    ("Domain alerts", "Low renewable and high net-grid thresholds", "Highlights operational energy conditions"),
    ("Airflow history", "Task/run state, retry and duration", "Diagnoses batch-file and billing failures"),
], [1750,3770,3840])
doc.add_heading("7.1 Alert Rules", level=2)
add_bullets(doc, ["LOW_RENEWABLE (WARNING): renewable_ratio < 0.10.", "HIGH_GRID_LOAD (CRITICAL): net_grid_usage_kwh > 8.0.", "NO RECENT DATA (health failure): last clean reading older than 120 seconds."])
doc.add_heading("7.2 Failure Demonstration", level=2)
doc.add_paragraph("Stopping meter-producer prevents new events. After 120 seconds, GET /health reports unhealthy and the last-event age explains why. Restarting the producer restores new Kafka events, Spark processing and a healthy response. Airflow failures retain audit error text and are retried; successful source files are skipped on future runs.")
add_picture(doc, "04-pipeline-alerts.png", "Figure 5. Dashboard alert evidence, including low-renewable warnings and high-grid-load critical events.", 6.1)

add_section_title(doc, "8", "Results and End-to-End Verification")
doc.add_heading("8.1 Representative Operational Result", level=2)
doc.add_paragraph("The captured dashboard reported 1,021.3 kWh consumption, 452.9 kWh solar generation, 568.4 kWh net grid load and 44.3% renewable contribution. EAST, NORTH, SOUTH and WEST were visible independently, allowing the operator to compare load and renewable mix by zone. These values are a time-specific demonstration snapshot and continue to change as simulated events arrive.")
doc.add_heading("8.2 Batch and Billing Result", level=2)
doc.add_paragraph("Airflow successfully processed the available tariff files and produced 30 current tariff rows and 30 household billing rows. The billing table exposes consumption, solar, billable grid usage, tariff, subsidy and estimated bill, making reconciliation traceable rather than presenting only a final total.")
add_picture(doc, "05-household-billing.png", "Figure 6. Consolidated household billing report produced from streaming usage and batch tariffs.", 6.15)
doc.add_heading("8.3 API Result", level=2)
doc.add_paragraph("GET /api/v1/overview returned HTTP 200 with total consumption, solar, net grid load, renewable ratio, reading count, last-updated timestamp and active-alert count. The /health endpoint reported healthy while events were current.")

add_section_title(doc, "9", "Orchestration and API Evidence")
doc.add_heading("9.1 Airflow Workflow", level=2)
doc.add_paragraph("The smart_grid_daily_tariff_pipeline DAG runs every five minutes. Its three ordered tasks discover files, load validated tariffs and generate the billing report. The captured Airflow view shows zero failed tasks/runs and repeated successful executions, demonstrating scheduling, orchestration and operational visibility.")
add_picture(doc, "07-airflow-successful-run.png", "Figure 7. Airflow grid view showing successful discover, load and billing tasks.", 6.2)
doc.add_heading("9.2 Serving Interface", level=2)
doc.add_paragraph("Swagger/OpenAPI documentation allows each endpoint to be inspected and executed during the demo. The overview response provides direct evidence that the serving layer reads consolidated data successfully.")
add_picture(doc, "09-api-overview-response.png", "Figure 8. Successful FastAPI overview request and JSON response.", 6.0)

add_section_title(doc, "10", "Testing, Limitations and Production Evolution")
doc.add_heading("10.1 Verification Performed", level=2)
add_table(doc, ["Check", "Observed result"], [
    ("Python syntax", "All application modules compiled successfully."),
    ("Unit tests", "Three business-logic tests passed."),
    ("Compose validation", "docker compose config --quiet passed."),
    ("Container status", "Eight application services were running; Kafka and PostgreSQL healthy."),
    ("Streaming", "Spark micro_batch_completed logs showed valid records and zero invalid in samples."),
    ("Batch", "Airflow successful; 30 tariffs and 30 billing records produced."),
    ("Serving", "Dashboard, API docs and Airflow UI returned HTTP 200; API health was healthy."),
], [2900,6460])
doc.add_heading("10.2 Current Limitations", level=2)
add_bullets(doc, ["A single Kafka broker, local Spark and one PostgreSQL instance provide no high availability.", "Spark writes each micro-batch through the driver; higher throughput needs a distributed or optimised sink.", "Kafka retention and raw-history policy are intentionally small for a laptop demo.", "Alerts are stored and displayed but are not routed to email, Slack or an on-call platform.", "Secrets and internal traffic use development defaults.", "The billing equation is illustrative and does not model real stepped utility tariffs or taxes."])
doc.add_heading("10.3 Production Evolution", level=2)
doc.add_paragraph("A production design would use replicated Kafka, a clustered Spark deployment, a schema registry, object storage for immutable raw history, managed PostgreSQL with replicas/backups, TLS and role-based access, a secrets manager, centralised logs/metrics/traces, and routed alerts. A scalable sink or lakehouse table format would replace driver-side writes. CI/CD would run tests, vulnerability scans and controlled schema migrations before deployment.")

add_section_title(doc, "11", "Conclusion")
doc.add_paragraph("The project delivers the required end-to-end data platform for Smart Grid Energy Monitoring & Billing. A continuous Python source feeds Kafka, Spark performs meaningful real-time validation and aggregation, and PostgreSQL stores queryable operational results. A second Python source produces simulated-daily tariffs that Airflow validates and reconciles with household usage. FastAPI and Streamlit expose the consolidated result, while structured logs, heartbeats, metrics, alerts and workflow history make the system observable.")
doc.add_paragraph("The Kappa-oriented decision is appropriate for this use case because low-latency meter insights are central, Kafka provides a replayable event log, and one processing implementation reduces consistency and maintenance risk. The retained Airflow path reflects the batch nature of external tariff reference data without duplicating meter-event business logic. The demonstrated system therefore answers both business questions: operators can view grid load and renewable contribution by zone, and billing staff can inspect an effective-tariff household report.")
doc.add_heading("References", level=1)
refs = [
    "EC8203 Applied Big Data Engineering. (2026). Data Engineering Mini-Project assessment brief.",
    "Apache Kafka Documentation. https://kafka.apache.org/documentation/",
    "Apache Spark Structured Streaming Programming Guide. https://spark.apache.org/docs/latest/structured-streaming-programming-guide.html",
    "Apache Airflow Documentation. https://airflow.apache.org/docs/",
    "PostgreSQL Documentation. https://www.postgresql.org/docs/",
    "FastAPI Documentation. https://fastapi.tiangolo.com/",
    "Streamlit Documentation. https://docs.streamlit.io/",
    "Docker Compose Documentation. https://docs.docker.com/compose/",
]
for ref in refs:
    p=doc.add_paragraph(style="List Number"); p.add_run(ref)
doc.add_heading("Submission Checklist", level=2)
add_bullets(doc, ["Replace the cover-page student/group placeholders.", "Review generated values and captions against the final demo run.", "Export this Word report to PDF for submission.", "Submit the repository/ZIP and prepare a 5-10 minute demo video or live demonstration.", "If submitted as a group, add an individual-contribution statement."])

doc.core_properties.title = "Smart Grid Energy Monitoring & Billing - Technical Report"
doc.core_properties.subject = "EC8203 Applied Big Data Engineering Mini Project"
doc.core_properties.author = "[Enter student name]"
doc.core_properties.keywords = "Kafka, Spark, Airflow, PostgreSQL, Kappa, Smart Grid"
doc.save(DOCX)
print(DOCX)
